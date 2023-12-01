import json
from fpdf import FPDF
from docx2pdf import convert
from io import BytesIO
from docx import Document
from docx.shared import Pt
import tempfile
import pythoncom
import streamlit as st


def load_cancer_info(cancer_info_json_path, cancer_type):
    with open(cancer_info_json_path, 'r') as file:
        cancer_info_list = json.load(file)

    for cancer_info in cancer_info_list:
        if cancer_info['type'] == cancer_type:
            return cancer_info
        
    return None


def write_self_care(cancer_info):
    for i in range(len(cancer_info['self_care']['heading'])):
        if cancer_info['self_care']['heading'][i] != '':
            st.subheader(cancer_info['self_care']['heading'][i].replace('\n', '  \n  \n'))
        st.write(cancer_info['self_care']['paragraph'][i].replace('\n', '  \n  \n'))


def write_resources(cancer_info):
    for i in range(len(cancer_info['resources']['paragraph'])):
        st.write(cancer_info['resources']['paragraph'][i].replace('\n', '  \n  \n'))
        st.link_button('Find out more', cancer_info['resources']['links'][i].replace('\n', '  \n  \n'))
        st.write('  \n')


def word_add_paragraph(document, text, font_size, font_type='Arial'):
    paragraph = document.add_paragraph(text)
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.name = font_type

def download_word(cancer_info):
    font_size_dict = {
        'title': 36,
        'heading': 24,
        'subheading': 16,
        'text':12
    }
    document = Document()
    
    word_add_paragraph(document, 'Cancer Vaccine App', font_size_dict['title'])

    word_add_paragraph(document, 'About your Cancer', font_size_dict['heading'])

    word_add_paragraph(document, cancer_info['about'], font_size_dict['text'])

    word_add_paragraph(document, 'Self-care', font_size_dict['heading'])

    for i in range(len(cancer_info['self_care']['heading'])):
        if cancer_info['self_care']['heading'][i] != '':
            word_add_paragraph(document, cancer_info['self_care']['heading'][i], font_size_dict['subheading'])

        word_add_paragraph(document, cancer_info['self_care']['paragraph'][i], font_size_dict['text'])

    if cancer_info['resources'] != '':
        word_add_paragraph(document, 'Resources', font_size_dict['heading'])

        for i in range(len(cancer_info['resources']['paragraph'])):
            word_add_paragraph(document, cancer_info['resources']['paragraph'][i], font_size_dict['text'])
            word_add_paragraph(document, "Link for more details: "+cancer_info['resources']['links'][i], font_size_dict['text'])

    # Save the document to BytesIO
    word_output = BytesIO()
    document.save(word_output)
    word_output.seek(0)

    # Save the Word document to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_word_file:
        temp_word_file_path = temp_word_file.name
        document.save(temp_word_file_path)

    # Convert Word to PDF using docx2pdf
    pythoncom.CoInitialize()  # Initialize COM
    try:
        temp_pdf_file_path = convert(temp_word_file_path)
    finally:
        pythoncom.CoUninitialize()  # Uninitialize COM

     # Construct the PDF file path
    temp_pdf_file_path = temp_word_file_path.replace(".docx", ".pdf")

    # Read the resulting PDF into a BytesIO object
    with open(temp_pdf_file_path, "rb") as pdf_file:
        pdf_output = BytesIO(pdf_file.read())


    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            label="DOCX",
            data=word_output,
            key="word_document",
            file_name="Cancer_info.docx",
            mime="application/octet-stream"
        )

    with col2:
        st.download_button(
            label="PDF",
            data=pdf_output,
            key="pdf_document",
            file_name="Cancer_info.pdf",
            mime="application/pdf"
        )

def render_consumer():

    cancer_info = load_cancer_info("cancer_info/cancer_info.json", "Breast")

    st.markdown("<h2 style='text-align: center;\
                    text-decoration: underline;\
                    padding-bottom: 7%;'>About your Cancer</h2>",
                    unsafe_allow_html=True)
    st.write(cancer_info['about'].replace('\n', '  \n  \n'))
    st.write('---')


    st.markdown("<h2 style='text-align: left;\
                    text-decoration: underline;\
                    padding-bottom: 7%;'>Self Care</h2>",
                    unsafe_allow_html=True)
    
    write_self_care(cancer_info)

    st.write('---')

    if cancer_info['resources'] != '':
        st.markdown("<h2 style='text-align: left;\
                        text-decoration: underline;\
                        padding-bottom: 7%;'>Resources</h2>",
                        unsafe_allow_html=True)
        
        write_resources(cancer_info)

    st.write('---')

    with st.container():
            subcol1, subcol2, subcol3 = st.columns([2,1,1])
            with subcol1:
               st.markdown("<h5 style='padding-top: 2%; text-align: right;'>Download as: </h2>",
                           unsafe_allow_html=True)
            with subcol2:
                download_word(cancer_info)
            with subcol3:
               pass
